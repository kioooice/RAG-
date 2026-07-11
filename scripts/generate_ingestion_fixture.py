"""Generate fictional MX-100 files; Office XLSX/PPTX are built by the companion artifact script."""
from pathlib import Path
import argparse,csv,json,shutil
from docx import Document
from docx.enum.text import WD_BREAK
from PIL import Image,ImageDraw,ImageFont
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

TEXT="MX-100 设备维护说明 v1.2\n温度上限 75℃；故障代码 E07。\n维护流程：1. 断电；2. 清洁滤网；3. 重启并记录结果。\n警告：温度达到75℃时立即停机。\n日期：2026-07-11 / Maintenance Guide"
def main():
 p=argparse.ArgumentParser();p.add_argument("--output-dir",type=Path,required=True);a=p.parse_args();d=a.output_dir;d.mkdir(parents=True,exist_ok=True)
 doc=Document();doc.add_heading("MX-100 设备维护说明",0);doc.add_paragraph("版本 v1.2 | 日期 2026-07-11");doc.add_heading("安全参数",1);doc.add_paragraph("设备型号 MX-100，温度上限 75℃，故障代码 E07。");doc.add_heading("维护流程",1)
 for s in ["断电","清洁滤网","重启并记录结果"]: doc.add_paragraph(s,style="List Number")
 doc.add_paragraph("警告：温度达到75℃时立即停机。");t=doc.add_table(rows=1,cols=3);t.rows[0].cells[0].text="代码";t.rows[0].cells[1].text="现象";t.rows[0].cells[2].text="处理"
 for row in [("E07","温度过高","停止设备并检查散热"),("E02","传感器离线","检查连接")]: cells=t.add_row().cells
 for i,v in enumerate(row): cells[i].text=v
 doc.save(d/"mx100_manual.docx")
 font_path=Path("C:/Windows/Fonts/msyh.ttc");pdfmetrics.registerFont(TTFont("CJK",str(font_path)));c=canvas.Canvas(str(d/"mx100_text.pdf"));c.setFont("CJK",16);y=790
 for line in TEXT.splitlines(): c.drawString(60,y,line);y-=32
 c.drawString(60,y,"代码 | 现象 | 处理");c.drawString(60,y-30,"E07 | 温度过高 | 检查散热");c.save()
 im=Image.new("RGB",(1400,900),"white");dr=ImageDraw.Draw(im);f=ImageFont.truetype(str(font_path),36);dr.multiline_text((60,60),TEXT+"\nE07 | 温度过高 | 检查散热",font=f,fill="black",spacing=18);im.save(d/"mx100_image.png")
 c=canvas.Canvas(str(d/"mx100_scanned.pdf"),pagesize=(700,450));c.drawImage(str(d/"mx100_image.png"),0,0,width=700,height=450);c.save()
 with (d/"mx100_devices.csv").open("w",encoding="utf-8-sig",newline="") as f: csv.writer(f).writerows([["model","version","max_temp","fault","date"],["MX-100","v1.2","75℃","E07","2026-07-11"]])
 (d/"mx100_config.json").write_text(json.dumps({"model":"MX-100","version":"v1.2","max_temperature":"75℃","fault_code":"E07","steps":["断电","清洁滤网","重启并记录结果"],"warning":"温度达到75℃时立即停机","date":"2026-07-11"},ensure_ascii=False,indent=2),encoding="utf-8")
 md="# MX-100 设备说明\n\n版本 v1.2；温度上限 75℃；故障代码 E07。\n\n1. 断电\n2. 清洁滤网\n3. 重启并记录结果\n\n> 警告：温度达到75℃时立即停机。\n\n日期：2026-07-11\n";(d/"mx100_notes.md").write_text(md,encoding="utf-8");shutil.copyfile(d/"mx100_notes.md",d/"mx100_notes_duplicate.md");c=canvas.Canvas(str(d/"wrong_extension.docx"));c.setFont("CJK",16);c.drawString(60,790,"MX-100 错误扩展名测试 v1.2 / E07 / 75℃");c.save();(d/"unsupported.bin").write_bytes(b"MX100-UNSUPPORTED-\x00\x01")
 gt_path=Path(__file__).resolve().parents[1]/"tests"/"fixtures"/"ingestion_ground_truth.json";(d.parent/"ground_truth.json").write_text(gt_path.read_text(encoding="utf-8"),encoding="utf-8")
if __name__=="__main__":main()
