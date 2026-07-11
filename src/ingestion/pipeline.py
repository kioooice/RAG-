"""Lightweight multi-format ingestion with deterministic IDs and traceability."""
from __future__ import annotations
import csv, hashlib, json, re, zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

RULE_VERSION = "ingestion-rules-v0.1"
NS = {"a":"http://schemas.openxmlformats.org/drawingml/2006/main","p":"http://schemas.openxmlformats.org/presentationml/2006/main"}

def norm(text: str) -> str:
    return "\n".join(" ".join(line.split()) for line in text.replace("\r\n","\n").replace("\r","\n").split("\n") if line.strip())
def sha(data: bytes) -> str: return hashlib.sha256(data).hexdigest()
def unit(doc_id, kind, title, text, locator, structured=None, metadata=None, parser="builtin"):
    clean=norm(text); key=json.dumps([doc_id,kind,title,locator,clean,structured],ensure_ascii=False,sort_keys=True).encode()
    return {"unit_id":"unit_"+sha(key)[:20],"document_id":doc_id,"unit_type":kind,"title":title,"section_path":[],"text":clean,"structured_data":structured,"source_locator":locator,"metadata":metadata or {},"quality_score":1.0 if clean or structured else 0.0,"parser_name":parser,"rule_version":RULE_VERSION}

def detect(path: Path):
    b=path.read_bytes(); ext=path.suffix.lower(); actual="unsupported"
    if b.startswith(b"%PDF"): actual="pdf"
    elif b.startswith(b"\x89PNG\r\n\x1a\n"): actual="png"
    elif b.startswith(b"PK\x03\x04"):
        with zipfile.ZipFile(path) as z:
            names=set(z.namelist())
        actual="docx" if any(n.startswith("word/") for n in names) else "xlsx" if any(n.startswith("xl/") for n in names) else "pptx" if any(n.startswith("ppt/") for n in names) else "unsupported"
    else:
        try:
            text=b.decode("utf-8-sig")
            if ext==".json": json.loads(text); actual="json"
            elif ext==".csv": list(csv.reader(text.splitlines())); actual="csv"
            elif ext in {".md",".markdown"}: actual="markdown"
        except Exception: pass
    expected={".docx":"docx",".pdf":"pdf",".xlsx":"xlsx",".pptx":"pptx",".png":"png",".csv":"csv",".json":"json",".md":"markdown"}.get(ext)
    return actual, bool(expected and expected!=actual)

def parse_docx(path, doc_id):
    d=DocxDocument(path); out=[]; section=[]; pi=0
    for p in d.paragraphs:
        t=norm(p.text)
        if not t: continue
        pi+=1
        if p.style and p.style.name.startswith("Heading"): section=[t]
        out.append(unit(doc_id,"paragraph",section[-1] if section else path.stem,t,{"paragraph":pi,"heading_path":section.copy()},metadata={"style":p.style.name if p.style else None},parser="python-docx"))
    for i,t in enumerate(d.tables,1):
        rows=[[norm(c.text) for c in r.cells] for r in t.rows]
        out.append(unit(doc_id,"table",f"Table {i}","",{"table":i},rows,parser="python-docx"))
    return out
def parse_pdf(path, doc_id):
    reader=PdfReader(path); out=[]
    for i,p in enumerate(reader.pages,1):
        text=norm(p.extract_text() or "")
        if text: out.append(unit(doc_id,"page",f"Page {i}",text,{"page":i},parser="pypdf"))
    return out, len(reader.pages)
def parse_xlsx(path, doc_id):
    wb=load_workbook(path,data_only=False,read_only=True); out=[]
    for ws in wb.worksheets:
        rows=[[c.value for c in row] for row in ws.iter_rows()]
        text="\n".join(" | ".join("" if v is None else str(v) for v in r) for r in rows)
        out.append(unit(doc_id,"table",ws.title,text,{"sheet":ws.title,"cell_range":ws.calculate_dimension(force=True)},rows,parser="openpyxl"))
    return out
def parse_pptx(path, doc_id):
    out=[]
    with zipfile.ZipFile(path) as z:
        slides=sorted([n for n in z.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml",n)],key=lambda n:int(re.search(r"\d+",n).group()))
        for i,name in enumerate(slides,1):
            root=ET.fromstring(z.read(name)); texts=[e.text or "" for e in root.findall(".//a:t",NS)]
            tables=[]
            for tbl in root.findall(".//a:tbl",NS):
                tables.append([[" ".join((t.text or "") for t in cell.findall(".//a:t",NS)) for cell in row.findall("a:tc",NS)] for row in tbl.findall("a:tr",NS)])
            out.append(unit(doc_id,"slide",f"Slide {i}","\n".join(texts),{"slide":i},tables or None,parser="openxml-pptx"))
    return out
def parse_csv_file(path, doc_id):
    with path.open(encoding="utf-8-sig",newline="") as f: rows=list(csv.reader(f))
    return [unit(doc_id,"table",path.stem,"\n".join(" | ".join(r) for r in rows),{"rows":f"1:{len(rows)}"},rows,parser="csv")]
def parse_json_file(path, doc_id):
    obj=json.loads(path.read_text(encoding="utf-8-sig")); return [unit(doc_id,"record",path.stem,json.dumps(obj,ensure_ascii=False),{"json_path":"$"},obj,parser="json")]
def parse_markdown(path, doc_id): return [unit(doc_id,"document",path.stem,path.read_text(encoding="utf-8-sig"),{"file":path.name},parser="markdown")]

def process_directory(input_dir: Path):
    seen_file={}; seen_content={}; docs=[]; units=[]
    for path in sorted(input_dir.iterdir(),key=lambda p:p.name.lower()):
        if not path.is_file(): continue
        raw=path.read_bytes(); checksum=sha(raw); actual,mismatch=detect(path); warnings=[]
        doc_id="doc_"+checksum[:20]
        if checksum in seen_file:
            docs.append({"document_id":doc_id,"source_path":str(path),"original_filename":path.name,"detected_file_type":actual,"file_size":len(raw),"checksum":checksum,"version":"v1.2","parser_name":"dedup","parser_version":"1","processing_status":"duplicate","quality_score":1.0,"warnings":[f"duplicate_of:{seen_file[checksum]}"],"processed_at":"deterministic"}); continue
        seen_file[checksum]=path.name; parsed=[]; status="accepted"; parser=actual
        try:
            if actual=="docx": parsed=parse_docx(path,doc_id); parser="python-docx"
            elif actual=="pdf":
                parsed,pages=parse_pdf(path,doc_id); parser="pypdf"
                if not parsed: status="needs_ocr"; warnings.append(f"no_extractable_text:{pages}_pages")
            elif actual=="xlsx": parsed=parse_xlsx(path,doc_id); parser="openpyxl"
            elif actual=="pptx": parsed=parse_pptx(path,doc_id); parser="openxml-pptx"
            elif actual=="csv": parsed=parse_csv_file(path,doc_id)
            elif actual=="json": parsed=parse_json_file(path,doc_id)
            elif actual=="markdown": parsed=parse_markdown(path,doc_id)
            elif actual=="png": status="needs_ocr"; warnings.append("image_has_no_machine_readable_text_layer")
            else: status="unsupported"; warnings.append("unsupported_file_type")
            if mismatch: status="needs_review"; warnings.append("extension_content_mismatch")
        except Exception as e: status="failed"; warnings.append(f"{type(e).__name__}:{e}")
        content_hash=sha("\n".join(u["text"]+json.dumps(u["structured_data"],ensure_ascii=False,sort_keys=True) for u in parsed).encode()) if parsed else None
        if content_hash and content_hash in seen_content: status="duplicate"; warnings.append(f"content_duplicate_of:{seen_content[content_hash]}"); parsed=[]
        elif content_hash: seen_content[content_hash]=path.name
        quality=1.0 if status=="accepted" else .8 if status in {"accepted_with_warnings","needs_review"} else .4 if status=="needs_ocr" else 0.0
        docs.append({"document_id":doc_id,"source_path":str(path),"original_filename":path.name,"detected_file_type":actual,"file_size":len(raw),"checksum":checksum,"version":"v1.2","parser_name":parser,"parser_version":"1","processing_status":status,"quality_score":quality,"warnings":warnings,"processed_at":"deterministic"}); units.extend(parsed)
    return docs,units
